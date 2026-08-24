"""Doc harvest: read the documentation found in a dataset folder (README / .md / .txt / .html / .docx / .pdf)
and attach description hints to every data file and field of a scan.

Two grades of hint:
  * ``auto``  - a line shaped like ``<field> | description`` / ``<field> : description`` / ``<field> - description``
                (schema tables in READMEs) -> can be used as the Description directly (flagged for review).
  * ``mention`` - any sentence that mentions the file/field name -> shown to the agent/user as context only.
"""
from __future__ import annotations

import html as _html
import re
from pathlib import Path
from typing import Optional

MAX_DOC_CHARS = 400_000
MAX_PDF_PAGES = 40
MAX_SNIPPETS = 3
SNIPPET_LEN = 260


def read_doc_text(path: Path) -> Optional[str]:
    ext = path.suffix.lower()
    try:
        if ext in (".md", ".txt", ".log", ".csvt"):
            return path.read_text(encoding="utf-8", errors="replace")[:MAX_DOC_CHARS]
        if ext in (".html", ".htm"):
            t = path.read_text(encoding="utf-8", errors="replace")
            t = re.sub(r"<script.*?</script>|<style.*?</style>", " ", t, flags=re.S | re.I)
            t = re.sub(r"<br\s*/?>|</p>|</tr>|</li>|</h\d>", "\n", t, flags=re.I)
            t = re.sub(r"</t[dh]>", " | ", t, flags=re.I)
            t = re.sub(r"<[^>]+>", " ", t)
            return _html.unescape(t)[:MAX_DOC_CHARS]
        if ext == ".docx":
            import docx  # python-docx (optional)
            d = docx.Document(str(path))
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for r in t.rows:
                    parts.append(" | ".join(c.text.strip() for c in r.cells))
            return "\n".join(parts)[:MAX_DOC_CHARS]
        if ext == ".pdf":
            import pdfplumber  # optional
            out = []
            with pdfplumber.open(str(path)) as pdf:
                for i, pg in enumerate(pdf.pages):
                    if i >= MAX_PDF_PAGES:
                        break
                    out.append(pg.extract_text() or "")
                    for tb in pg.extract_tables() or []:
                        for r in tb:
                            out.append(" | ".join((c or "").replace("\n", " ") for c in r))
            return "\n".join(out)[:MAX_DOC_CHARS]
    except Exception:
        return None
    return None


def _sentences(text: str) -> list[str]:
    lines = []
    for ln in text.splitlines():
        ln = ln.strip()
        if not ln or set(ln) <= set("-=|*_#:` "):
            continue
        lines.append(ln)
    return lines


def _clean(s: str) -> str:
    s = re.sub(r"`|\*{1,3}", "", s).strip(" |:-–—\t")   # underscores kept: snake_case identifiers are common
    return re.sub(r"\s+", " ", s)[:SNIPPET_LEN]


def harvest(folder: Path, scan: dict) -> dict:
    """Mutates *scan*: adds scan['docs'] (per document: chars, lines) and per data entry
    e['doc_hints'] = {'file': {...}, 'fields': {field: {...}}}. Returns a summary."""
    docs: dict[str, list[str]] = {}
    for name in list(scan.get("documents", [])) + [e["name"] for e in scan["files"] if e.get("format") == "TXT"]:
        p = folder / name
        if not p.exists():
            continue
        txt = read_doc_text(p)
        if txt:
            docs[name] = _sentences(txt)
    scan["docs"] = {k: {"lines": len(v)} for k, v in docs.items()}
    n_auto = n_mention = 0
    for e in scan["files"]:
        if e["role"] != "data":
            continue
        targets = {"__file__": Path(e["name"]).name}
        fields = (e.get("fields") or [])
        if not fields and e.get("inner"):
            continue
        for c in fields:
            targets[c["name"]] = c["name"]
        hints: dict = {"file": {"auto": None, "mentions": []}, "fields": {}}
        stem = re.escape(Path(e["name"]).name)
        stem_noext = re.escape(Path(e["name"]).name.split(".")[0])
        file_rx = re.compile(rf"(?<![\w]){stem}(?![\w])|(?<![\w]){stem_noext}(?![\w])", re.I)
        for doc, lines in docs.items():
            for i, ln in enumerate(lines):
                if file_rx.search(ln) and len(hints["file"]["mentions"]) < MAX_SNIPPETS:
                    hints["file"]["mentions"].append({"doc": doc, "text": _clean(ln)})
                    # a heading followed by a purpose line (README style)
                    if hints["file"]["auto"] is None and i + 1 < len(lines) and re.match(r"^\**(purpose|description|תיאור|מטרה)", lines[i + 1], re.I):
                        hints["file"]["auto"] = {"doc": doc, "text": _clean(re.sub(r"^\**(purpose|description|תיאור|מטרה)\**\s*[:\-–]\s*", "", lines[i + 1], flags=re.I))}
                        n_auto += 1
        for fname in targets:
            if fname == "__file__":
                continue
            rx_auto = re.compile(rf"^\s*[`*|]*\s*{re.escape(fname)}\s*[`*]*\s*(\||:|-|–|—|\t)\s*(.+?)\s*\|?\s*$", re.I)
            rx_ment = re.compile(rf"(?<![\w]){re.escape(fname)}(?![\w])", re.I)
            fh = {"auto": None, "mentions": []}
            for doc, lines in docs.items():
                for ln in lines:
                    m = rx_auto.match(ln)
                    if m and fh["auto"] is None and len(m.group(2)) > 2 and not re.fullmatch(r"[\w\s]*", m.group(2)) or (m and fh["auto"] is None and len(m.group(2).split()) >= 2):
                        fh["auto"] = {"doc": doc, "text": _clean(m.group(2))}
                        n_auto += 1
                    elif rx_ment.search(ln) and len(fh["mentions"]) < MAX_SNIPPETS and not m:
                        fh["mentions"].append({"doc": doc, "text": _clean(ln)})
                        n_mention += 1
            if fh["auto"] or fh["mentions"]:
                hints["fields"][fname] = fh
        if hints["file"]["auto"] or hints["file"]["mentions"] or hints["fields"]:
            e["doc_hints"] = hints
    summary = {"documents_read": len(docs), "auto_descriptions": n_auto, "mentions": n_mention}
    scan["doc_harvest"] = summary
    return summary
